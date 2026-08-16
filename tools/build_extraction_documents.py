import json
import re
import subprocess
from collections import defaultdict
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SOURCE = Path("build/extractions/source")
OUTPUT = Path("docs/extractions")
ACCENT = RGBColor(91, 63, 224)
INK = RGBColor(32, 36, 43)
MUTED = RGBColor(95, 99, 110)
LIGHT = "F3F1FF"

EDITORIAL = {
    "574cd3c8-d550-4e0f-a671-04b016434131": {
        "title": "Alur Produksi dan Strategi Maintenance",
        "summary": "Rekaman membahas gambaran alur material dari penerimaan dan penimbangan hingga compounding, filling, dan packaging, lalu beralih ke pendekatan maintenance yang memadukan breakdown, preventive, time-based, dan predictive maintenance.",
        "themes": [
            "Alur produksi: penimbangan, tagging, penyimpanan, compounding, filling, dan packaging.",
            "Perbedaan kemasan primer, sekunder, dan tersier serta hubungannya dengan mutu produk.",
            "Pemilihan strategi maintenance berdasarkan tingkat kritikal komponen.",
            "Kondisi saat ini masih didominasi breakdown dan preventive maintenance; predictive maintenance memerlukan data dan instrumen yang lebih matang.",
        ],
        "actions": [
            "Petakan mesin dan komponen berdasarkan dampaknya terhadap keselamatan, kualitas, dan kontinuitas produksi.",
            "Tentukan komponen yang cukup run-to-failure, yang wajib diganti berkala, dan yang layak dipantau secara prediktif.",
            "Lengkapi alur produksi dengan pemilik proses, titik kontrol mutu, dan catatan handoff antarbagian.",
        ],
    },
    "3f5fabe9-953c-483e-93ec-68004fb9210c": {
        "title": "Observasi Mesin, Maintenance, dan Audit Operasional",
        "summary": "Rekaman panjang ini menangkap observasi lapangan dan diskusi teknis mengenai kondisi mesin, pisau dan setting, gangguan listrik atau suplai, aktivitas perawatan, pengujian, keselamatan, serta kesiapan menghadapi audit. Percakapan bersifat sangat spontan sehingga beberapa istilah teknis perlu dikonfirmasi kembali.",
        "themes": [
            "Pemeriksaan fungsi mesin, termasuk kemampuan memotong, setting, penggantian pisau, dan kondisi saat mesin berhenti.",
            "Gangguan listrik, suplai, kontrol, dan kebutuhan penelusuran penyebab masalah.",
            "Pengujian, post-test, pelaporan masalah, dan tindak lanjut setelah perbaikan.",
            "Keselamatan kerja dan frekuensi audit internal maupun eksternal.",
        ],
        "actions": [
            "Validasi nama mesin, kode aset, dan istilah teknis yang tertangkap kurang jelas sebelum dijadikan laporan resmi.",
            "Ubah temuan lapangan menjadi daftar masalah dengan status, pemilik, akar penyebab, dan target penyelesaian.",
            "Satukan bukti pengujian dan tindakan koreksi sebagai paket kesiapan audit.",
        ],
    },
    "938b357b-807b-446e-83c1-5a65ea789285": {
        "title": "Building Maintenance, Project, dan Peran Engineering",
        "summary": "Rekaman membahas cakupan building maintenance untuk mendukung produksi dan kepatuhan, termasuk bangunan, AC, pencahayaan, dan aset fasilitas. Diskusi juga menyentuh rotasi atau pembelajaran lintas fungsi, project engineering, otomasi, microcontroller, AI, serta ekspektasi kompetensi peserta.",
        "themes": [
            "Building maintenance tidak terbatas pada struktur bangunan, tetapi mencakup utilitas dan aset pendukung operasional.",
            "Tujuan utama adalah menjaga kelancaran produksi sekaligus memenuhi regulasi.",
            "Project engineering dan peran PIC memberi paparan lintas disiplin namun membutuhkan kepemilikan yang jelas.",
            "Otomasi, microcontroller, machine learning, dan AI muncul sebagai area pembelajaran yang relevan.",
        ],
        "actions": [
            "Susun daftar aset building dan batas tanggung jawab antara building, utility, workshop, dan project.",
            "Buat matriks kompetensi peserta untuk maintenance, fasilitas, otomasi, dan project management.",
            "Tentukan proyek kecil yang terukur agar pembelajaran teknis menghasilkan keluaran nyata.",
        ],
    },
    "dbcca064-95c4-4d1c-bd66-285c289de4f1": {
        "title": "Pengolahan Air Limbah dan Pengelolaan Limbah B3",
        "summary": "Rekaman mendokumentasikan tur proses pengolahan air limbah: penetralan pH dan suhu, equalisasi, tahap anaerob dan aerasi, pengendalian bakteri, clarifier, lumpur aktif, hingga penanganan sludge dan limbah B3. Sejumlah nama alat atau bahan terdengar tidak pasti dan perlu verifikasi teknis.",
        "themes": [
            "Kontrol pH dan suhu sebelum limbah memasuki proses biologis.",
            "Peran proses anaerob dan aerasi dalam menurunkan beban organik seperti COD dan BOD.",
            "Pemantauan kondisi bakteri dan lumpur aktif sebagai indikator kesehatan proses.",
            "Penanganan sludge, proses press, serta pemisahan dan pengelolaan limbah B3.",
        ],
        "actions": [
            "Konfirmasi batas operasi pH, suhu, COD, BOD, dan parameter lumpur dengan SOP laboratorium.",
            "Buat diagram alir resmi dari titik masuk limbah sampai keluaran dan pengangkutan sludge.",
            "Verifikasi istilah alat, bahan kimia, dan nama limbah yang akurasinya rendah pada transkrip.",
        ],
    },
    "8ac94d57-53a4-4c22-80f3-9a80037b0a2d": {
        "title": "Sistem Kelistrikan, Genset, dan Proteksi Kebakaran",
        "summary": "Rekaman berisi penjelasan lapangan tentang komponen kelistrikan, perbedaan AC dan DC, tegangan satu dan tiga fasa, transformator, genset, panel, serta sistem pompa kebakaran. Aspek bahaya tegangan tinggi dan urutan kerja jockey, electric, dan diesel pump menjadi bagian penting.",
        "themes": [
            "Pengenalan kontaktor, relay, ACB, sistem AC/DC, dan level tegangan umum.",
            "Transformasi tegangan utilitas menuju kebutuhan 380 V dan 220 V.",
            "Peran genset dan panel dalam menjaga kontinuitas daya.",
            "Logika tekanan pada jockey pump, electric pump, dan diesel pump untuk proteksi kebakaran.",
        ],
        "actions": [
            "Pastikan seluruh angka tegangan, kapasitas, dan set point tekanan diverifikasi terhadap single-line diagram dan SOP.",
            "Tambahkan penandaan bahaya dan batas akses di area tegangan tinggi.",
            "Dokumentasikan urutan start pompa kebakaran dan jadwal uji berkala.",
        ],
    },
    "76b931df-2c8b-434f-95b3-a373e5a5a71b": {
        "title": "Kerangka Wawancara dan Evaluasi Proses",
        "summary": "Rekaman singkat membahas cara menggali proses melalui sasaran, pemasok dan konsumen, relasi internal, aktivitas rutin, dampak positif, tantangan, serta pelajaran. Percakapan mengarah pada kerangka wawancara yang dapat dipakai untuk memetakan proses dan pengembangan individu.",
        "themes": [
            "Tujuan proses dan indikator sasaran.",
            "Hubungan pemasok, konsumen, dan pihak internal yang mendukung proses.",
            "Aktivitas rutin, laporan, evaluasi, dan koordinasi.",
            "Dampak positif, tantangan, pelajaran, dan aspirasi pengembangan.",
        ],
        "actions": [
            "Gunakan urutan pertanyaan yang konsisten: tujuan, progres, relasi eksternal, relasi internal, hal positif, tantangan, dan pelajaran.",
            "Minta contoh dan bukti untuk setiap jawaban agar hasil tidak berhenti pada persepsi.",
            "Pisahkan temuan proses dari refleksi pengembangan personal dalam laporan akhir.",
        ],
    },
}


def clean(text):
    return re.sub(r"\s+", " ", text or "").strip()


def clock(seconds):
    seconds = int(max(0, seconds))
    return f"{seconds // 3600:02d}:{(seconds % 3600) // 60:02d}:{seconds % 60:02d}"


def set_font(run, name="Aptos", size=None, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def shade(paragraph, fill=LIGHT):
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def configure(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, before, after in [
        ("Title", 28, 0, 8), ("Subtitle", 13, 0, 18),
        ("Heading 1", 16, 16, 8), ("Heading 2", 13, 12, 6), ("Heading 3", 11.5, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Aptos Display" if name in ("Title", "Heading 1") else "Aptos"
        style._element.rPr.rFonts.set(qn("w:ascii"), style.font.name)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), style.font.name)
        style.font.size = Pt(size)
        style.font.color.rgb = ACCENT if name.startswith("Heading") else INK
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    bullet = doc.styles["List Bullet"]
    bullet.font.name = "Aptos"
    bullet.font.size = Pt(10.5)
    bullet.paragraph_format.left_indent = Inches(0.5)
    bullet.paragraph_format.first_line_indent = Inches(-0.25)
    bullet.paragraph_format.space_after = Pt(6)
    bullet.paragraph_format.line_spacing = 1.1

    if "Transcript Timestamp" not in doc.styles:
        ts = doc.styles.add_style("Transcript Timestamp", WD_STYLE_TYPE.PARAGRAPH)
        ts.base_style = normal
        ts.font.name = "Aptos"
        ts.font.size = Pt(9.2)
        ts.font.color.rgb = INK
        ts.paragraph_format.space_after = Pt(5)
        ts.paragraph_format.line_spacing = 1.05
        ts.paragraph_format.keep_together = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run("VOXTRACE  /  EXTRACTION REPORT")
    set_font(run, size=8, bold=True, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Generated from stored VoxTrace extraction")
    set_font(run, size=8, color=MUTED)


def add_cover(doc, title, subtitle, meta_lines):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(72)
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("VOXTRACE / TRANSCRIPT INTELLIGENCE")
    set_font(r, size=9, bold=True, color=ACCENT)
    p = doc.add_paragraph(title, style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p = doc.add_paragraph(subtitle, style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(18)
    shade(p)
    r = p.add_run("\n".join(meta_lines))
    set_font(r, size=10.5, color=INK)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(80)
    r = p.add_run("Catatan kualitas\n")
    set_font(r, size=9, bold=True, color=MUTED)
    r = p.add_run("Dokumen ini merapikan hasil speech-to-text otomatis. Istilah, nama, angka, dan bagian dengan audio tumpang tindih tetap perlu diverifikasi terhadap rekaman sumber sebelum dipakai sebagai bukti formal.")
    set_font(r, size=9, color=MUTED)
    doc.add_page_break()


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_timeline(doc, segments, bucket_seconds=600):
    buckets = defaultdict(list)
    for seg in segments:
        buckets[int(seg["start"] // bucket_seconds)].append(seg)
    for bucket, items in sorted(buckets.items()):
        start = bucket * bucket_seconds
        end = min((bucket + 1) * bucket_seconds, max(x["end"] for x in items))
        p = doc.add_paragraph(style="Heading 2")
        p.add_run(f"{clock(start)} - {clock(end)}")
        text = clean(" ".join(x["text"] for x in items))
        preview = text[:480].rsplit(" ", 1)[0]
        if len(text) > len(preview):
            preview += "..."
        doc.add_paragraph(preview)


def add_full_transcript(doc, segments, window_seconds=120):
    groups = defaultdict(list)
    for seg in segments:
        groups[int(seg["start"] // window_seconds)].append(seg)
    for _, items in sorted(groups.items()):
        start, end = items[0]["start"], items[-1]["end"]
        speakers = {clean(x.get("speaker", "")) for x in items if clean(x.get("speaker", ""))}
        label = ", ".join(sorted(speakers)) if speakers and speakers != {"SPEAKER_00"} else "Ucapan"
        p = doc.add_paragraph(style="Transcript Timestamp")
        r = p.add_run(f"{clock(start)} - {clock(end)}  |  {label}\n")
        set_font(r, size=8.5, bold=True, color=ACCENT)
        r = p.add_run(clean(" ".join(x["text"] for x in items)))
        set_font(r, size=9.2, color=INK)


def metadata_for(job_id):
    query = (
        "SELECT json_build_object('model',j.model,'createdAt',j.created_at,'completedAt',j.completed_at,"
        "'processing',pm.metadata)::text FROM jobs j LEFT JOIN processing_metadata pm ON pm.job_id=j.id "
        f"WHERE j.id='{job_id}';"
    )
    result = subprocess.run(
        ["docker", "compose", "exec", "-T", "postgres", "psql", "-U", "voxtrace", "-d", "voxtrace", "-At", "-c", query],
        check=True, capture_output=True, text=True, encoding="utf-8",
    )
    return json.loads(result.stdout.strip())


def build_record(record):
    info = EDITORIAL[record["id"]]
    meta = metadata_for(record["id"])
    proc = meta.get("processing") or {}
    doc = Document()
    configure(doc)
    add_cover(doc, info["title"], record["filename"], [
        f"Durasi audio: {clock(record['duration'])}",
        f"Bahasa terdeteksi: {record['language']}",
        f"Segmen sumber: {len(record['segments']):,}",
        f"Backend: {proc.get('backend', 'unknown')} / model {proc.get('model', meta.get('model', 'unknown'))}",
        f"Dibuat: {datetime.now().strftime('%d %B %Y')}",
    ])
    doc.add_heading("Ringkasan eksekutif", level=1)
    p = doc.add_paragraph(info["summary"])
    p.paragraph_format.keep_together = True
    doc.add_heading("Pokok bahasan", level=1)
    add_bullets(doc, info["themes"])
    doc.add_heading("Tindak lanjut yang disarankan", level=1)
    add_bullets(doc, info["actions"])
    doc.add_heading("Peta isi berdasarkan waktu", level=1)
    doc.add_paragraph("Cuplikan berikut membantu navigasi; transkrip penuh tersedia setelah bagian ini.")
    add_timeline(doc, record["segments"])
    transcript_heading = doc.add_heading("Transkrip tertata", level=1)
    transcript_heading.paragraph_format.page_break_before = True
    doc.add_paragraph("Teks dikelompokkan dalam jendela sekitar dua menit dan tetap mempertahankan urutan kronologis hasil ekstraksi.")
    add_full_transcript(doc, record["segments"])
    target = OUTPUT / f"{record['filename'].rsplit('.', 1)[0]} - Laporan Ekstraksi.docx"
    doc.save(target)
    return target, meta


def build_master(records):
    doc = Document()
    configure(doc)
    total = sum(x["duration"] for x in records)
    add_cover(doc, "Ikhtisar Hasil Ekstraksi", "Kumpulan rekaman operasional dan pembelajaran lapangan", [
        f"Rekaman selesai: {len(records)}",
        f"Total durasi: {clock(total)}",
        f"Total segmen: {sum(len(x['segments']) for x in records):,}",
        "Bahasa utama: Indonesia",
        f"Dibuat: {datetime.now().strftime('%d %B %Y')}",
    ])
    doc.add_heading("Gambaran umum", level=1)
    doc.add_paragraph("Kumpulan rekaman berfokus pada orientasi proses produksi dan engineering: alur produksi, strategi maintenance, utilitas listrik, proteksi kebakaran, building maintenance, pengolahan air limbah, serta kerangka evaluasi proses. Hasil ini paling berguna sebagai bahan awal knowledge capture dan harus dikonfirmasi oleh pemilik proses untuk istilah teknis, nama aset, angka, dan persyaratan regulasi.")
    doc.add_heading("Tema lintas rekaman", level=1)
    add_bullets(doc, [
        "Keandalan operasi bergantung pada pemetaan aset kritis dan pemilihan strategi maintenance yang sesuai.",
        "Keselamatan, utilitas, dan kepatuhan perlu dibuktikan melalui SOP, inspeksi, pengujian, serta catatan audit.",
        "Pengetahuan lapangan tersebar dalam percakapan spontan; standardisasi istilah dan penetapan pemilik proses akan meningkatkan keterpakaiannya.",
        "Peluang perbaikan terbesar adalah mengubah observasi menjadi register tindakan yang memiliki status, PIC, tenggat, dan bukti penutupan.",
    ])
    doc.add_heading("Katalog dokumen", level=1)
    for index, record in enumerate(records, 1):
        info = EDITORIAL[record["id"]]
        p = doc.add_paragraph(style="Heading 2")
        p.add_run(f"{index}. {info['title']}")
        p = doc.add_paragraph()
        r = p.add_run(f"{record['filename']} | {clock(record['duration'])} | {len(record['segments']):,} segmen\n")
        set_font(r, size=9, bold=True, color=ACCENT)
        p.add_run(info["summary"])
    doc.add_heading("Prioritas verifikasi", level=1)
    add_bullets(doc, [
        "Nama mesin, kode aset, bahan kimia, dan istilah lokal yang terdengar ambigu.",
        "Set point, kapasitas, tegangan, tekanan, frekuensi inspeksi, serta angka kepatuhan.",
        "Pernyataan yang akan dijadikan dasar keputusan, audit, atau perubahan prosedur.",
        "Identitas pembicara pada rekaman tanpa diarization atau dengan audio yang saling tumpang tindih.",
    ])
    doc.add_heading("Paket keluaran", level=1)
    doc.add_paragraph("Setiap rekaman memiliki laporan mandiri yang berisi ringkasan eksekutif, tema utama, rekomendasi tindak lanjut, peta waktu, dan transkrip lengkap yang telah dikelompokkan agar mudah dibaca.")
    target = OUTPUT / "00 - Ikhtisar Hasil Ekstraksi VoxTrace.docx"
    doc.save(target)
    return target


def main():
    OUTPUT.mkdir(parents=True, exist_ok=True)
    records = []
    for path in SOURCE.glob("*.json"):
        data = json.loads(path.read_text(encoding="utf-8"))
        if isinstance(data, dict) and data.get("id") in EDITORIAL and data.get("segments"):
            records.append(data)
    records.sort(key=lambda x: x["filename"])
    outputs = [build_master(records)]
    for record in records:
        target, _ = build_record(record)
        outputs.append(target)
    print("\n".join(str(x) for x in outputs))


if __name__ == "__main__":
    main()

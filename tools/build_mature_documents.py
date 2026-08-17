from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = Path("docs/extractions")
PURPLE = RGBColor(91, 63, 224)
INK = RGBColor(28, 31, 36)
MUTED = RGBColor(92, 97, 107)

DOCS = [
    {
        "file": "13-08-2026 10.36 - Panduan Alur Produksi dan Maintenance.docx",
        "title": "Alur Produksi dan Strategi Maintenance",
        "subtitle": "Panduan operasional untuk memahami aliran material, titik kendali, dan pemilihan pola pemeliharaan mesin produksi",
        "purpose": "Memberikan gambaran utuh mengenai hubungan alur produksi dengan keandalan mesin, sehingga gangguan dapat dicegah, ditangani, dan dievaluasi berdasarkan dampaknya terhadap mutu serta kontinuitas produksi.",
        "scope": ["Penerimaan dan penimbangan bahan", "Penyimpanan dan identifikasi material", "Compounding, filling, dan packaging", "Klasifikasi kemasan primer, sekunder, dan tersier", "Breakdown, preventive, time-based, dan predictive maintenance"],
        "principles": [
            ("Aliran terkendali", "Setiap perpindahan material memiliki identitas, status, pemilik, dan titik serah yang jelas."),
            ("Mutu melekat pada proses", "Kontrol tidak hanya dilakukan pada produk akhir; parameter penting dijaga sejak penimbangan hingga pengemasan."),
            ("Strategi berbasis kritikalitas", "Pola maintenance dipilih menurut risiko keselamatan, kualitas, penghentian produksi, biaya, dan waktu pemulihan."),
            ("Bukti sebelum penutupan", "Pekerjaan dinyatakan selesai setelah fungsi diuji dan hasilnya tercatat."),
        ],
        "process": [
            ("1", "Penerimaan & penimbangan", "Verifikasi bahan, jumlah, identitas, dan status pelepasan."),
            ("2", "Tagging & penyimpanan", "Jaga keterlacakan bahan dan pisahkan status material secara visual maupun sistem."),
            ("3", "Compounding", "Proses pencampuran mengikuti formula dan parameter tervalidasi."),
            ("4", "Filling", "Produk diisikan ke kemasan primer; kesiapan mesin dan kebersihan menjadi titik kendali utama."),
            ("5", "Packaging", "Kemasan sekunder/tersier melindungi, mengidentifikasi, dan menyiapkan produk untuk distribusi."),
            ("6", "Rekonsiliasi & rilis", "Bandingkan pemakaian, hasil, reject, dan sisa; selesaikan penyimpangan sebelum penutupan batch."),
        ],
        "controls": [
            ("Ketersediaan mesin", "Downtime, frekuensi gangguan, dan waktu pemulihan", "Harian/mingguan"),
            ("Kualitas proses", "Parameter kritis dan reject pada tiap tahap", "Per batch"),
            ("Preventive maintenance", "Kepatuhan jadwal dan pekerjaan tertunda", "Mingguan/bulanan"),
            ("Keandalan komponen", "Pola kerusakan dan kebutuhan suku cadang", "Bulanan"),
        ],
        "risks": [
            ("Kegagalan komponen kritis", "Henti produksi atau risiko mutu", "Daftar kritikalitas, preventive maintenance, dan stok minimum komponen kritis"),
            ("Serah-terima tidak lengkap", "Material atau status proses sulit ditelusuri", "Checklist handoff dan identifikasi status"),
            ("Perbaikan tanpa uji fungsi", "Gangguan berulang", "Post-test terdokumentasi sebelum work order ditutup"),
        ],
        "actions": ["Susun peta mesin dan komponen kritis per tahapan produksi.", "Tetapkan pemilik proses serta bukti serah-terima pada setiap handoff.", "Pisahkan komponen run-to-failure, penggantian berkala, dan kandidat predictive maintenance.", "Gunakan data gangguan untuk memperbarui jadwal serta kebutuhan suku cadang."],
        "notes": ["Catatan pendukung menguatkan urutan compounding–filling–packaging dan peran mesin penunjang produksi.", "Angka parameter proses spesifik harus mengikuti dokumen produksi dan validasi resmi; tidak ditetapkan oleh panduan ini."],
    },
    {
        "file": "12-08-2026 09.05 - Panduan Observasi Mesin dan Suku Cadang.docx",
        "title": "Observasi Mesin, Maintenance, dan Tata Kelola Suku Cadang",
        "subtitle": "Panduan lapangan untuk mengubah temuan teknis menjadi tindakan korektif yang terlacak dan siap diaudit",
        "purpose": "Menyatukan pola observasi mesin, penanganan gangguan, pembuktian perbaikan, serta permintaan dan pengeluaran suku cadang tanpa mengambil alih kewenangan purchasing atau pengelola inventori.",
        "scope": ["Observasi fungsi dan kondisi mesin", "Diagnosis awal dan eskalasi gangguan", "Work request dan work order", "Permintaan, pengeluaran, serta bukti pemakaian suku cadang", "Post-test, penutupan pekerjaan, dan kesiapan audit"],
        "principles": [
            ("Fakta sebelum diagnosis", "Catat gejala, kondisi operasi, waktu, dan bukti; bedakan observasi dari dugaan penyebab."),
            ("Satu jejak pekerjaan", "Work request, work order, material, tindakan, pengujian, dan penutupan saling merujuk."),
            ("Pemisahan kewenangan", "Pengguna menjelaskan kebutuhan; spare-part/inventory mengendalikan stok; purchasing menjalankan pengadaan."),
            ("Bukti pemakaian", "Material yang dikeluarkan dapat ditelusuri ke pekerjaan dan aset yang menerima material."),
        ],
        "process": [
            ("1", "Temukan & amankan", "Hentikan atau batasi operasi bila ada risiko keselamatan, kualitas, atau kerusakan lanjutan."),
            ("2", "Buat work request", "Tuliskan aset, gejala, kondisi, prioritas, pelapor, waktu, dan bukti pendukung."),
            ("3", "Rencanakan pekerjaan", "Tentukan tindakan, tenaga, izin, alat, material, dan kebutuhan downtime."),
            ("4", "Cek ketersediaan", "Gunakan stok yang sesuai; bila tidak tersedia, ajukan permintaan pengadaan melalui alur persetujuan."),
            ("5", "Release & eksekusi WO", "Catat material/consumable yang dikeluarkan terhadap work order lalu lakukan pekerjaan."),
            ("6", "Post-test & penutupan", "Uji fungsi, rekam hasil, dokumentasikan material terpakai/sisa, dan tutup WO di sistem."),
        ],
        "controls": [
            ("Respons gangguan", "Waktu laporan, mulai penanganan, dan pemulihan", "Per kejadian"),
            ("Kelengkapan WO", "Aset, tindakan, material, hasil uji, dan PIC", "Sebelum closing"),
            ("Ketersediaan suku cadang", "Min–max, pemakaian, lead time, dan shortage", "Mingguan/bulanan"),
            ("Kesiapan audit", "Keterhubungan bukti dari permintaan sampai pemakaian", "Bulanan"),
        ],
        "risks": [
            ("Permintaan tidak terpenuhi", "Waktu perbaikan memanjang", "Kritikalitas stok, min–max, alternatif teknis yang disetujui, dan eskalasi shortage"),
            ("Material keluar tanpa referensi", "Stok dan biaya tidak dapat ditelusuri", "Wajib tautkan release material ke WO dan aset"),
            ("Diagnosis berubah menjadi fakta", "Tindakan salah atau catatan audit lemah", "Pisahkan gejala, hipotesis, pengujian, dan akar penyebab terkonfirmasi"),
        ],
        "actions": ["Gunakan format temuan: aset–gejala–dampak–bukti–tindakan–hasil.", "Tetapkan critical spare dan parameter min–max berdasarkan konsumsi, lead time, serta dampak downtime.", "Pastikan PR hanya diterbitkan setelah kebutuhan teknis dan anggaran disetujui; PO tetap menjadi kewenangan purchasing.", "Lakukan rekonsiliasi material terpakai, dikembalikan, dan rusak sebelum WO ditutup."],
        "notes": ["Catatan pendukung memperjelas alur WR → input material/consumable → release → WO → pencatatan sistem (disebut Oracle dalam catatan).", "Nilai target belanja, nama formulir, nama pemberi persetujuan, dan angka kekurangan 5–10% masih perlu dikonfirmasi sebelum menjadi data resmi."],
    },
    {
        "file": "12-08-2026 13.27 - Panduan Building Maintenance dan Project.docx",
        "title": "Building Maintenance dan Tata Kelola Project Engineering",
        "subtitle": "Kerangka pengelolaan fasilitas untuk menjaga kepatuhan, keselamatan, dan kelancaran operasi",
        "purpose": "Menetapkan batas kerja building maintenance, hubungan dengan utility/workshop/project, serta cara mengendalikan pekerjaan operasional dan proyek tanpa mencampur kewenangan antarunit.",
        "scope": ["Bangunan: lantai, dinding, kaca, atap, pintu, dan fasilitas umum", "AC/HVAC pada sisi fasilitas", "Pencahayaan dan instalasi pendukung", "IPAL, TPS, dan fasilitas keselamatan pada sisi sarana", "Pekerjaan operasional berbasis WO dan pekerjaan proyek"],
        "principles": [
            ("Batas tanggung jawab jelas", "Building menjaga sarana; utility menjaga pasokan; workshop menangani mesin; project mengelola perubahan terencana."),
            ("Produksi dan regulasi", "Prioritas pekerjaan mempertimbangkan kelancaran produksi sekaligus kepatuhan."),
            ("Perubahan terkendali", "Project memiliki ruang lingkup, PIC, anggaran, risiko, persetujuan, dan kriteria selesai."),
            ("Kinerja dapat dibuktikan", "WO closing mengukur operasi; project done mengukur penyelesaian proyek, masing-masing dengan definisi yang konsisten."),
        ],
        "process": [
            ("1", "Identifikasi kebutuhan", "Pisahkan gangguan fasilitas, pekerjaan rutin, perbaikan, dan perubahan proyek."),
            ("2", "Klasifikasi & prioritas", "Nilai keselamatan, kepatuhan, mutu, dampak produksi, biaya, dan urgensi."),
            ("3", "Penetapan pemilik", "Tentukan building, utility, workshop, project, atau kolaborasi dengan satu PIC utama."),
            ("4", "Perencanaan", "Susun metode, izin, material, vendor, waktu kerja, dan mitigasi gangguan operasi."),
            ("5", "Eksekusi & pengawasan", "Pastikan keselamatan, kualitas pekerjaan, serta pengendalian perubahan."),
            ("6", "Verifikasi & serah-terima", "Uji fungsi, lengkapi as-built/dokumen, perbarui daftar aset, lalu tutup pekerjaan."),
        ],
        "controls": [
            ("Pekerjaan operasional", "WO closing tepat waktu dan backlog berdasarkan risiko", "Mingguan/bulanan"),
            ("Pekerjaan proyek", "Project done terhadap ruang lingkup, waktu, biaya, dan acceptance", "Per milestone"),
            ("Kondisi fasilitas", "Temuan inspeksi dan pengulangan kerusakan", "Rutin"),
            ("Keselamatan", "Insiden, near miss, dan temuan kepatuhan", "Per kejadian/bulanan"),
        ],
        "risks": [
            ("Area abu-abu antarfungsi", "Pekerjaan tertunda atau saling lempar", "Matriks RACI dan satu PIC untuk setiap pekerjaan"),
            ("Project selesai tanpa serah-terima", "Aset sulit dirawat", "Acceptance, manual, drawing, daftar spare, dan pelatihan"),
            ("Target persentase tidak terdefinisi", "Kinerja sulit dibandingkan", "Definisikan numerator, denominator, periode, dan sumber data"),
        ],
        "actions": ["Bangun register aset fasilitas beserta pemilik dan batas layanan.", "Buat matriks prioritas berdasarkan keselamatan, regulasi, mutu, dan dampak produksi.", "Tetapkan definisi formal WO closing dan project done sebelum memakai target persentase.", "Gunakan proyek kecil otomasi/monitoring sebagai sarana belajar bila ruang lingkup dan manfaatnya terukur."],
        "notes": ["Catatan menyebut capaian sekitar 90,08% (2025) dan 94,22% (2026), namun indikator dan periodenya belum terverifikasi sehingga tidak diperlakukan sebagai KPI resmi.", "Istilah MTTR/MTBF dicoret pada catatan; dokumen ini tidak menetapkannya sebagai KPI building maintenance."],
    },
    {
        "file": "12-08-2026 14.28 - Panduan IPAL dan Limbah B3.docx",
        "title": "Pengendalian IPAL, Sludge, dan Limbah B3",
        "subtitle": "Panduan proses, titik kendali, dan bukti operasional pengolahan air limbah",
        "purpose": "Menjelaskan aliran pengolahan limbah cair dan pengelolaan residunya agar operator dapat memahami tujuan setiap tahap, parameter yang harus dijaga, serta bukti yang diperlukan untuk kepatuhan.",
        "scope": ["Penerimaan dan equalisasi limbah", "Koreksi pH dan suhu", "Proses anaerob dan aerasi", "Clarifier dan pengendalian lumpur aktif", "Dewatering/sludge press dan pengelolaan limbah B3"],
        "principles": [
            ("Stabilkan sebelum proses biologis", "Fluktuasi pH, suhu, dan beban diserap di tahap awal agar mikroorganisme tidak mengalami shock."),
            ("Bakteri adalah aset proses", "Aerasi, nutrisi, beban, dan karakter lumpur dipantau sebagai satu sistem."),
            ("Pisahkan air dan residu", "Clarifier mengendapkan biomassa; sludge berlebih ditangani melalui jalur yang terdokumentasi."),
            ("Data laboratorium mengendalikan keputusan", "Perubahan operasi mengikuti hasil ukur tervalidasi, tren, dan SOP—notasi catatan lapangan bukan batas izin."),
        ],
        "process": [
            ("1", "Inlet & screening/equalisasi", "Tampung dan homogenkan aliran; catat volume dan karakter awal."),
            ("2", "Kondisioning", "Sesuaikan pH dan suhu sesuai rentang operasi yang disahkan."),
            ("3", "Anaerob", "Kurangi sebagian beban organik dalam kondisi tanpa oksigen terlarut sesuai desain."),
            ("4", "Aerasi", "Pasok udara dan jaga populasi mikroorganisme untuk menguraikan beban organik."),
            ("5", "Clarifier", "Pisahkan air olahan dan lumpur; kembalikan lumpur aktif sesuai kebutuhan proses."),
            ("6", "Outlet & sludge", "Verifikasi kualitas efluen; press, simpan, label, dan serahkan residu melalui jalur B3 yang sah."),
        ],
        "controls": [
            ("Kondisi inlet", "Debit/volume, pH, suhu, dan indikasi beban", "Per shift/harian"),
            ("Proses biologis", "Aerasi, kondisi lumpur, pH, suhu, dan tren proses", "Per shift/harian"),
            ("Kualitas outlet", "Parameter sesuai baku mutu dan metode laboratorium", "Sesuai izin/SOP"),
            ("Sludge/B3", "Volume, label, penyimpanan, manifest, dan serah-terima", "Setiap perpindahan"),
        ],
        "risks": [
            ("Shock load", "Proses biologis terganggu", "Equalisasi, isolasi aliran abnormal, dan eskalasi cepat"),
            ("Interpretasi angka tidak tervalidasi", "Keputusan operasi atau kepatuhan salah", "Gunakan SOP laboratorium, izin, dan hasil uji resmi"),
            ("Sludge tidak terlacak", "Risiko lingkungan dan audit", "Neraca massa sederhana dan dokumen pengangkutan lengkap"),
        ],
        "actions": ["Buat diagram proses resmi lengkap dengan arah aliran, recycle, sampling point, dan emergency diversion.", "Tetapkan dashboard tren inlet–proses–outlet, bukan hanya angka sesaat.", "Hubungkan hasil uji dengan tindakan koreksi dan penanggung jawab.", "Rekonsiliasi volume air olahan, sludge, dan limbah B3 secara periodik."],
        "notes": ["Catatan menampilkan angka COD 100, BOD 30, dan TSS <10 serta volume 146/170/130 m³; seluruh angka diperlakukan sebagai referensi wawancara yang wajib dicocokkan dengan izin, SOP, desain, dan hasil laboratorium.", "Angka volume tahunan/biaya yang tertulis tidak cukup jelas untuk diterbitkan sebagai fakta dan sengaja tidak dimasukkan sebagai target."],
    },
    {
        "file": "12-08-2026 11.22 - Panduan Kelistrikan dan Proteksi Kebakaran.docx",
        "title": "Sistem Kelistrikan, Daya Darurat, dan Proteksi Kebakaran",
        "subtitle": "Panduan orientasi untuk memahami arsitektur sistem, urutan fungsi, dan batas keselamatan kerja",
        "purpose": "Memberikan kerangka aman untuk memahami distribusi daya, genset, panel, dan rangkaian pompa kebakaran. Dokumen ini bersifat orientasi dan tidak menggantikan single-line diagram, SOP switching, atau prosedur tanggap darurat.",
        "scope": ["Konsep AC/DC dan sistem satu/tiga fasa", "Transformator, panel, dan proteksi", "Genset dan kontinuitas daya", "Fire alarm, hydrant, serta rangkaian pompa kebakaran", "Inspeksi, pengujian, dan pembatasan akses"],
        "principles": [
            ("Energi harus dikendalikan", "Identifikasi sumber, isolasi, verifikasi bebas energi, grounding, dan izin kerja mendahului pekerjaan."),
            ("Diagram adalah sumber kebenaran", "Nomor panel, kapasitas, tegangan, interlock, dan jalur suplai diverifikasi pada dokumen resmi."),
            ("Redundansi harus diuji", "Genset dan pompa cadangan bernilai hanya bila start, transfer, dan beban dapat dibuktikan."),
            ("Proteksi berlapis", "Fire alarm mendeteksi; hydrant menyalurkan; jockey menjaga tekanan; electric/diesel pump menyediakan debit utama sesuai desain."),
        ],
        "process": [
            ("1", "Pemetaan sumber", "Kenali sumber utilitas, transformator, genset, panel utama, dan beban kritis."),
            ("2", "Distribusi & proteksi", "Ikuti jalur panel dan perangkat proteksi tanpa mengandalkan ingatan atau label informal."),
            ("3", "Daya darurat", "Uji logika start, transfer, alarm, bahan bakar, dan kemampuan membawa beban yang ditetapkan."),
            ("4", "Deteksi kebakaran", "Pastikan alarm, detector, panel, annunciation, dan jalur eskalasi bekerja."),
            ("5", "Pemadaman berbasis air", "Verifikasi reservoir, hydrant/sprinkler, valve, tekanan, serta urutan pompa."),
            ("6", "Pemulihan & laporan", "Kembalikan konfigurasi normal, catat deviasi, dan tindak lanjuti temuan."),
        ],
        "controls": [
            ("Kelistrikan", "Kondisi panel, proteksi, thermography/inspeksi, dan housekeeping", "Sesuai PM"),
            ("Genset", "Start, transfer, alarm, bahan bakar, baterai, dan uji beban", "Sesuai SOP"),
            ("Fire alarm", "Deteksi, notifikasi, dan komunikasi", "Sesuai regulasi/SOP"),
            ("Fire pump", "Tekanan, auto-start, valve, kebocoran, dan urutan operasi", "Sesuai regulasi/SOP"),
        ],
        "risks": [
            ("Angka tegangan/set point salah", "Cedera atau sistem gagal bekerja", "Verifikasi pada diagram, nameplate, dan SOP sebelum digunakan"),
            ("Pengujian tanpa koordinasi", "Gangguan produksi atau alarm tak terkendali", "Permit, pemberitahuan, pembatasan area, dan rencana pemulihan"),
            ("Mode manual tertinggal", "Proteksi otomatis tidak tersedia", "Checklist pengembalian ke posisi normal"),
        ],
        "actions": ["Perbarui single-line diagram dan daftar beban kritis.", "Buat matriks sumber normal–darurat untuk setiap beban penting.", "Dokumentasikan urutan jockey–electric–diesel pump berdasarkan cause-and-effect resmi.", "Gunakan checklist pengujian yang memuat kondisi awal, hasil, deviasi, dan restorasi."],
        "notes": ["Catatan pendukung menegaskan peran operator pada hydrant dan fire alarm serta hubungan utility dengan suplai listrik.", "Nilai tegangan, kapasitas, dan set point tekanan dari percakapan tidak dijadikan instruksi kerja; semuanya wajib diverifikasi pada dokumen resmi."],
    },
    {
        "file": "13-08-2026 13.37 - Kerangka Evaluasi Proses.docx",
        "title": "Kerangka Wawancara dan Evaluasi Proses",
        "subtitle": "Panduan praktis menghasilkan profil proses yang berbasis bukti, konsisten, dan dapat ditindaklanjuti",
        "purpose": "Mengarahkan wawancara agar menghasilkan pemahaman proses yang utuh—tujuan, relasi, aktivitas, dampak, tantangan, pembelajaran, dan agility—tanpa mencampurkan isi dari unit yang sedang dievaluasi dengan unit lain.",
        "scope": ["Persiapan wawancara", "Tujuh ranah pertanyaan", "Penggalian contoh dan bukti", "Pemisahan fakta, interpretasi, dan aspirasi", "Sintesis serta tindak lanjut"],
        "principles": [
            ("Satu proses, satu konteks", "Informasi lintas unit dipakai untuk memperjelas antarmuka, bukan mengubah substansi pemilik proses."),
            ("Bukti menyertai pernyataan", "Setiap klaim penting ditopang contoh, dokumen, data, atau observasi."),
            ("Netral terhadap narasumber", "Pertanyaan tidak menggiring; kontradiksi dicatat untuk verifikasi, bukan diselesaikan dengan asumsi."),
            ("Tindakan memiliki pemilik", "Rekomendasi memuat hasil yang diharapkan, PIC, tenggat, dan bukti selesai."),
        ],
        "process": [
            ("1", "Purpose", "Apa mandat proses, pelanggan, keluaran, dan indikator keberhasilannya?"),
            ("2", "External relation", "Siapa pemasok/pelanggan eksternal dan apa komitmen layanan yang berlaku?"),
            ("3", "Internal relation", "Unit mana yang memberi input, menerima output, menyetujui, atau mendukung?"),
            ("4", "Positive impact", "Nilai apa yang tercipta dan bukti apa yang menunjukkan dampaknya?"),
            ("5", "Challenge", "Hambatan, risiko, bottleneck, dan ketergantungan apa yang paling menentukan?"),
            ("6", "Lesson learned", "Apa yang berubah setelah kejadian, audit, proyek, atau kegagalan sebelumnya?"),
            ("7", "Agility", "Seberapa cepat proses merespons perubahan kebutuhan, prioritas, dan kondisi lapangan?"),
        ],
        "controls": [
            ("Kelengkapan", "Tujuh ranah terjawab dengan contoh", "Setiap wawancara"),
            ("Kualitas bukti", "Proporsi klaim penting dengan referensi", "Saat review"),
            ("Konsistensi", "Istilah dan batas proses sama antarbagian", "Saat sintesis"),
            ("Tindak lanjut", "Temuan memiliki pemilik dan status", "Mingguan"),
        ],
        "risks": [
            ("Jawaban normatif", "Profil proses tidak menggambarkan praktik", "Tanyakan contoh terakhir, frekuensi, bukti, dan pengecualian"),
            ("Informasi lintas unit mengintervensi", "Dokumen kehilangan fokus", "Gunakan informasi lintas unit hanya pada bagian antarmuka dan beri atribusi"),
            ("Angka tanpa definisi", "Kesimpulan menyesatkan", "Minta formula, periode, sumber, pemilik, dan target pembanding"),
        ],
        "actions": ["Kirim agenda dan daftar bukti sebelum wawancara.", "Gunakan tujuh ranah secara konsisten tetapi izinkan pertanyaan lanjutan.", "Tandai setiap pernyataan sebagai terverifikasi, perlu verifikasi, atau interpretasi.", "Validasi ringkasan kepada pemilik proses sebelum diterbitkan."],
        "notes": ["Kerangka tujuh ranah berasal dari catatan pendukung dan diperjelas oleh rekaman singkat.", "Dokumen ini adalah metode evaluasi; contoh dari produksi, utility, building, IPAL, atau spare part tidak diperlakukan sebagai isi proses ini."],
    },
]

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), fill); tcPr.append(shd)

def set_cell_text(cell, text, bold=False, color=INK):
    cell.text = ""; p = cell.paragraphs[0]; r = p.add_run(str(text)); r.bold = bold; r.font.name = "Aptos"; r.font.size = Pt(9); r.font.color.rgb = color
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def setup(doc):
    sec = doc.sections[0]; sec.page_width = Inches(8.27); sec.page_height = Inches(11.69)
    sec.top_margin = Inches(.72); sec.bottom_margin = Inches(.68); sec.left_margin = sec.right_margin = Inches(.78)
    normal = doc.styles["Normal"]; normal.font.name = "Aptos"; normal.font.size = Pt(9.5); normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(5); normal.paragraph_format.line_spacing = 1.08
    for n, size in [("Title", 27), ("Subtitle", 12), ("Heading 1", 16), ("Heading 2", 11.5)]:
        s=doc.styles[n]; s.font.name="Aptos Display" if n in ("Title","Heading 1") else "Aptos"; s.font.size=Pt(size); s.font.color.rgb=PURPLE if n.startswith("Heading") else INK
        s.font.bold=n!="Subtitle"; s.paragraph_format.keep_with_next=True; s.paragraph_format.space_before=Pt(12); s.paragraph_format.space_after=Pt(6)
    h=sec.header.paragraphs[0]; h.text="VOXTRACE  /  OPERATIONAL KNOWLEDGE"; h.runs[0].font.size=Pt(8); h.runs[0].font.bold=True; h.runs[0].font.color.rgb=MUTED
    f=sec.footer.paragraphs[0]; f.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    f.add_run("Internal working document  •  17 August 2026  •  ")
    fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); f._p.append(fld)
    for r in f.runs: r.font.size=Pt(8); r.font.color.rgb=MUTED

def cover(doc, d):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(58); r=p.add_run("VOXTRACE / KNOWLEDGE BRIEF"); r.bold=True; r.font.size=Pt(9); r.font.color.rgb=PURPLE
    doc.add_paragraph(d["title"], "Title"); doc.add_paragraph(d["subtitle"], "Subtitle")
    t=doc.add_table(rows=2, cols=2); t.autofit=False; t.columns[0].width=Inches(2.0); t.columns[1].width=Inches(4.7)
    set_cell_text(t.cell(0,0),"STATUS",True,RGBColor(255,255,255)); shade(t.cell(0,0),"5B3FE0"); set_cell_text(t.cell(0,1),"Dokumen matang • sumber terkorelasi • angka kritis terkontrol")
    set_cell_text(t.cell(1,0),"POSISI DOKUMEN",True,RGBColor(255,255,255)); shade(t.cell(1,0),"22252B"); set_cell_text(t.cell(1,1),"Panduan domain mandiri; bukan transkrip dan bukan pengganti SOP resmi")
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(30); r=p.add_run("TUJUAN\n"); r.bold=True; r.font.color.rgb=PURPLE; p.add_run(d["purpose"])
    doc.add_paragraph("Ruang lingkup", "Heading 2")
    for x in d["scope"]: doc.add_paragraph(x, "List Bullet")
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(18); r=p.add_run("Prinsip korelasi sumber\n"); r.bold=True; r.font.color.rgb=PURPLE
    p.add_run("Rekaman adalah sumber utama domain. Catatan scan dan foto hanya memperjelas konteks yang relevan. Informasi dari domain lain tidak digunakan untuk mengubah mandat, angka, atau keputusan dokumen ini.")
    doc.add_page_break()

def table(doc, headers, rows, widths=None):
    t=doc.add_table(rows=1, cols=len(headers)); t.style="Table Grid"
    for i,h in enumerate(headers): set_cell_text(t.rows[0].cells[i],h,True,RGBColor(255,255,255)); shade(t.rows[0].cells[i],"5B3FE0")
    for ri,row in enumerate(rows):
        cells=t.add_row().cells
        for i,v in enumerate(row): set_cell_text(cells[i],v); shade(cells[i],"F5F3FF" if ri%2==0 else "FFFFFF")
    if widths:
        for row in t.rows:
            for i,w in enumerate(widths): row.cells[i].width=Inches(w)
    doc.add_paragraph()

def build_one(d):
    doc=Document(); setup(doc); cover(doc,d)
    doc.add_heading("1. Model kerja",1); table(doc,["Prinsip","Makna operasional"],d["principles"],[1.8,4.8])
    doc.add_heading("2. Alur kerja yang disarankan",1); table(doc,["Tahap","Aktivitas","Keluaran minimum"],d["process"],[.55,1.65,4.4])
    doc.add_heading("3. Kendali dan bukti",1); table(doc,["Area","Bukti/indikator","Frekuensi"],d["controls"],[1.5,3.7,1.4])
    doc.add_heading("4. Risiko utama dan pengendalian",1); table(doc,["Risiko","Dampak","Pengendalian"],d["risks"],[1.55,1.65,3.4])
    doc.add_heading("5. Agenda implementasi",1)
    for i,x in enumerate(d["actions"],1): doc.add_paragraph(f"{i}. {x}")
    doc.add_heading("6. Catatan penggunaan dan verifikasi",1)
    for x in d["notes"]: doc.add_paragraph(x,"List Bullet")
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(10); shade_para=OxmlElement("w:shd"); shade_para.set(qn("w:fill"),"FFF4D6"); p._p.get_or_add_pPr().append(shade_para)
    r=p.add_run("Batas penggunaan. "); r.bold=True; p.add_run("Dokumen ini mendukung orientasi, diskusi, dan penyusunan kontrol. Untuk pekerjaan berisiko, kepatuhan, parameter proses, atau switching, gunakan SOP/izin/drawing/hasil laboratorium yang berlaku.")
    doc.save(OUT/d["file"])

def build_index():
    doc=Document(); setup(doc)
    d={"title":"Portofolio Pengetahuan Operasional","subtitle":"Peta enam dokumen domain hasil pematangan rekaman dan catatan pendukung","purpose":"Menyediakan pintu masuk ke enam dokumen yang berdiri sendiri, sekaligus memperlihatkan titik antarmuka tanpa mencampur isi atau kewenangan tiap domain.","scope":["Produksi dan maintenance", "Observasi mesin dan suku cadang", "Building maintenance dan project", "IPAL dan limbah B3", "Kelistrikan dan proteksi kebakaran", "Evaluasi proses"]}
    cover(doc,d)
    doc.add_heading("1. Struktur portofolio",1)
    rows=[]
    for x in DOCS: rows.append((x["title"],x["purpose"],x["file"]))
    table(doc,["Dokumen","Mandat inti","Berkas"],rows,[1.7,3.15,1.75])
    doc.add_heading("2. Aturan korelasi",1)
    for x in ["Setiap dokumen memakai rekamannya sebagai sumber inti.","Catatan tambahan hanya masuk jika memperjelas konteks domain yang sama.","Fakta lintas domain ditempatkan sebagai antarmuka, bukan sebagai perubahan mandat.","Angka, istilah, nama, dan set point yang belum jelas diberi status perlu verifikasi.","SOP, izin, drawing, hasil laboratorium, dan data sistem tetap menjadi sumber formal tertinggi."]:
        doc.add_paragraph(x,"List Bullet")
    doc.add_heading("3. Antarmuka utama",1)
    table(doc,["Dari","Ke","Objek serah-terima"],[
        ("Produksi","Maintenance","Gejala, dampak proses, window pekerjaan, hasil post-test"),
        ("Maintenance","Spare part/Purchasing","Kebutuhan teknis, kritikalitas, WO/PR, bukti pemakaian"),
        ("Building","Utility/Project","Batas aset, risiko, izin, perubahan, dan acceptance"),
        ("IPAL","Laboratorium/EHS","Sampling, hasil uji, deviasi, sludge/B3, dan bukti kepatuhan"),
        ("Kelistrikan/Fire","Operasi/EHS","Rencana uji, isolasi, alarm, status proteksi, dan restorasi"),
        ("Evaluator","Pemilik proses","Klaim, bukti, status verifikasi, keputusan, dan tindak lanjut"),
    ],[1.45,1.45,3.6])
    doc.add_heading("4. Status kematangan",1)
    doc.add_paragraph("Seluruh dokumen telah diubah dari laporan ekstraksi menjadi panduan operasional. Transkrip mentah dan timeline tidak lagi menjadi tubuh dokumen. Ketidakpastian yang masih tersisa dikumpulkan pada bagian verifikasi di dokumen terkait sehingga tidak menyebar ke domain lain.")
    doc.save(OUT/"00 - Ikhtisar Pengetahuan Operasional VoxTrace.docx")

if __name__ == "__main__":
    OUT.mkdir(parents=True,exist_ok=True)
    for d in DOCS: build_one(d)
    build_index()
    print("Built",len(DOCS)+1,"mature documents")
